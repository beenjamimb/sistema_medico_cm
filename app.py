import io
import os
from datetime import datetime
import pandas as pd
import streamlit as st

# Manipulação de DOCX
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

st.set_page_config(
    page_title="Emissor Médico Especializado", page_icon="🩺", layout="centered"
)


# ---------------------------------------------------------
# CARREGAMENTO E CACHE DA BASE CID-10
# ---------------------------------------------------------
@st.cache_data
def carregar_biblioteca_cid10():
  for encoding in ["latin1", "utf-8", "iso-8859-1"]:
    for sep in [";", ","]:
      try:
        df_cat = pd.read_csv(
            "CID-10-CATEGORIAS.CSV", sep=sep, encoding=encoding, dtype=str
        )
        df_sub = pd.read_csv(
            "CID-10-SUBCATEGORIAS.CSV", sep=sep, encoding=encoding, dtype=str
        )

        df_cat.columns = [c.strip().lower() for c in df_cat.columns]
        df_sub.columns = [c.strip().lower() for c in df_sub.columns]

        col_cat_code = "cat" if "cat" in df_cat.columns else df_cat.columns[0]
        col_cat_desc = (
            "descricao" if "descricao" in df_cat.columns else df_cat.columns[1]
        )

        col_sub_code = (
            "subcat" if "subcat" in df_sub.columns else df_sub.columns[0]
        )
        col_sub_desc = (
            "descricao" if "descricao" in df_sub.columns else df_sub.columns[1]
        )

        df_cat["label"] = (
            df_cat[col_cat_code].astype(str)
            + " - "
            + df_cat[col_cat_desc].astype(str)
        )
        df_sub["label"] = (
            df_sub[col_sub_code].astype(str)
            + " - "
            + df_sub[col_sub_desc].astype(str)
        )

        df_todos = pd.concat(
            [df_sub[["label"]], df_cat[["label"]]], ignore_index=True
        ).drop_duplicates()
        return df_todos
      except Exception:
        continue

  st.warning(
      "Não foi possível carregar automaticamente os arquivos"
      " 'CID-10-CATEGORIAS.CSV' e 'CID-10-SUBCATEGORIAS.CSV'."
  )
  return pd.DataFrame(columns=["label"])


def set_cell_border(cell, **kwargs):
  """Auxiliar para definir bordas personalizadas nas células de tabela do Word."""
  tcPr = cell._tc.get_or_add_tcPr()
  tcBorders = tcPr.first_child_found_in("w:tcBorders")
  if tcBorders is None:
    tcBorders = OxmlElement("w:tcBorders")
    tcPr.append(tcBorders)
  for edge in ("top", "left", "bottom", "right"):
    edge_data = kwargs.get(edge)
    if edge_data:
      tag = "w:{}".format(edge)
      element = tcBorders.find(qn(tag))
      if element is None:
        element = OxmlElement(tag)
        tcBorders.append(element)
      for key, val in edge_data.items():
        element.set(qn("w:{}".format(key)), str(val))


def adicionar_receituario_especial(
    doc, nome_medico, nome_paciente, prescricao
):
  hoje = datetime.now()
  meses = [
      "Janeiro",
      "Fevereiro",
      "Março",
      "Abril",
      "Maio",
      "Junho",
      "Julho",
      "Agosto",
      "Setembro",
      "Outubro",
      "Novembro",
      "Dezembro",
  ]
  data_extenso_str = f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

  vias = [
      "1ª VIA - RETENÇÃO DA FARMÁCIA / DROGARIA",
      "2ª VIA - ORIENTAÇÃO AO PACIENTE",
  ]

  for idx, via in enumerate(vias):
    # --- TÍTULO ---
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_after = Pt(8)
    run_tit = p_tit.add_run(f"RECEITUÁRIO DE CONTROLE ESPECIAL\n({via})")
    run_tit.bold = True
    run_tit.font.name = "Arial"
    run_tit.font.size = Pt(11)

    # --- IDENTIFICAÇÃO DO EMITENTE ---
    p_emit = doc.add_paragraph()
    p_emit.paragraph_format.space_after = Pt(8)
    run_emit = p_emit.add_run(
        "IDENTIFICAÇÃO DO EMITENTE\n"
        f"Nome: {nome_medico}\n"
        "Endereço: Av. do Contorno, s/n - São Luís, Campo Maior - PI,"
        " 64280-000\n"
        "Telefone: (86) 3562-1192 | Cidade: Campo Maior  UF: PI"
    )
    run_emit.font.name = "Arial"
    run_emit.font.size = Pt(9)

    # --- PACIENTE ---
    p_pac = doc.add_paragraph()
    p_pac.paragraph_format.space_after = Pt(10)
    run_pac = p_pac.add_run(f"PACIENTE: {nome_paciente}")
    run_pac.bold = True
    run_pac.font.name = "Arial"
    run_pac.font.size = Pt(10.5)

    # --- PRESCRIÇÃO ---
    p_presc_tit = doc.add_paragraph()
    p_presc_tit.paragraph_format.space_after = Pt(4)
    run_pt = p_presc_tit.add_run("Prescrição:")
    run_pt.bold = True
    run_pt.font.name = "Arial"
    run_pt.font.size = Pt(10.5)

    for linha in prescricao.split("\n"):
      p_line = doc.add_paragraph()
      p_line.paragraph_format.space_after = Pt(2)
      p_line.paragraph_format.left_indent = Inches(0.2)
      run_line = p_line.add_run(linha)
      run_line.font.name = "Arial"
      run_line.font.size = Pt(10)

    # --- DATA E ASSINATURA DO MÉDICO ---
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_data.paragraph_format.space_before = Pt(15)
    p_data.paragraph_format.space_after = Pt(10)
    run_data = p_data.add_run(
        f"Campo Maior - PI, {data_extenso_str}\n\n_____________________________________\nAssinatura"
        " e Carimbo do Médico"
    )
    run_data.font.name = "Arial"
    run_data.font.size = Pt(9.5)

    # --- CAMPOS INFERIORES: COMPRADOR E FORNECEDOR ---
    table_rodape = doc.add_table(rows=1, cols=2)
    table_rodape.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_kw = {"val": "single", "sz": "4", "space": "0", "color": "CCCCCC"}

    # Célula 1: Identificação do Comprador
    cell_comp = table_rodape.cell(0, 0)
    set_cell_border(
        cell_comp,
        top=border_kw,
        bottom=border_kw,
        left=border_kw,
        right=border_kw,
    )
    p_comp = cell_comp.paragraphs[0]
    p_comp.paragraph_format.space_before = Pt(4)
    p_comp.paragraph_format.space_after = Pt(4)
    run_comp = p_comp.add_run(
        "IDENTIFICAÇÃO DO COMPRADOR\n"
        "Nome: _________________________________\n"
        "RG Nº: ________________ Órg. Emissor: ____\n"
        "Endereço: _____________________________\n"
        "Cidade: ____________________ UF: ______\n"
        "Telefone: ______________________________"
    )
    run_comp.font.name = "Arial"
    run_comp.font.size = Pt(8.5)

    # Célula 2: Identificação do Fornecedor
    cell_forn = table_rodape.cell(0, 1)
    set_cell_border(
        cell_forn,
        top=border_kw,
        bottom=border_kw,
        left=border_kw,
        right=border_kw,
    )
    p_forn = cell_forn.paragraphs[0]
    p_forn.paragraph_format.space_before = Pt(4)
    p_forn.paragraph_format.space_after = Pt(4)
    run_forn = p_forn.add_run(
        "IDENTIFICAÇÃO DO FORNECEDOR\n\n"
        "_______________________________________\n"
        "Assinatura do Farmacêutico / Data\n\n"
        "Data do Atendimento: ____ / ____ / ______"
    )
    run_forn.font.name = "Arial"
    run_forn.font.size = Pt(8.5)

    if idx == 0:
      doc.add_page_break()


# ---------------------------------------------------------
# INTERFACE PRINCIPAL
# ---------------------------------------------------------
if "pagina_atual" not in st.session_state:
  st.session_state["pagina_atual"] = "menu"

senha = st.text_input("Digite sua chave de acesso:", type="password")

if senha == "trabalheconquiste":

  if st.session_state["pagina_atual"] == "menu":
    st.markdown(
        "<h1 style='text-align: center;'>🩺 Sistema de Emissão de Documentos"
        " Médicos</h1>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<p style='text-align: center; color: gray;'>Selecione o módulo ou"
        " funcionalidade desejada</p>",
        unsafe_allow_html=True,
    )
    st.write("\n")

    col_m1, col_m2 = st.columns(2)
    with col_m1:
      if st.button(
          "🩺\n\nDocumentos Cirurgia Geral", use_container_width=True
      ):
        st.session_state["pagina_atual"] = "cirurgia_geral"
        st.rerun()

      st.button(
          "📝\n\nTermo de Consentimento Livre e Esclarecido - TCLE (em breve)",
          use_container_width=True,
          disabled=True,
      )

      st.button(
          "🏥\n\nPedido de Vaga de UTI (em breve)",
          use_container_width=True,
          disabled=True,
      )

    with col_m2:
      if st.button("🦴\n\nDocumentos Ortopedia", use_container_width=True):
        st.session_state["pagina_atual"] = "ortopedia"
        st.rerun()

      # Estilização CSS para centralizar os elementos do link_button (Emoji + Texto)
      st.markdown(
          """
          <style>
          div[data-testid="stLinkButton"] > a {
              display: flex;
              flex-direction: column;
              align-items: center;
              justify-content: center;
              text-align: center;
          }
          </style>
          """,
          unsafe_allow_html=True,
      )

      # Botão Beta de Hemocomponentes
      st.link_button(
          "🩸\n\nRequisição de Hemocomponentes (Beta)",
          "https://www.hemopi.pi.gov.br/formularios/requisicao-hemopi/",
          use_container_width=True,
      )
      st.caption("⚠️ Em desenvolvimento - imprimir e preencher a mão")

      st.button(
          "✂️\n\nDescrições Cirúrgicas (em breve)",
          use_container_width=True,
          disabled=True,
      )

  # ---------------------------------------------------------
  # CIRURGIA GERAL
  # ---------------------------------------------------------
  elif st.session_state["pagina_atual"] == "cirurgia_geral":
    if st.button("⬅️ Voltar ao Menu Principal"):
      st.session_state["pagina_atual"] = "menu"
      st.rerun()

    st.header("🩺 Emissor de Documentos - Cirurgia Geral")

    st.subheader("1. Selecione os Documentos a serem gerados")
    emitir_atestado = st.checkbox("Atestado Médico", value=False)
    emitir_retorno = st.checkbox(
        "Ficha de Retorno (Ambulatório)", value=False
    )
    emitir_receita = st.checkbox("Receituário Comum", value=False)
    emitir_receita_especial = st.checkbox(
        "Receituário de Controle Especial (2 Vias)", value=False
    )
    emitir_histo = st.checkbox("Solicitação de Histopatológico", value=False)
    emitir_antibio = st.checkbox("Solicitação de Antibiograma", value=False)

    algum_selecionado = (
        emitir_atestado
        or emitir_retorno
        or emitir_receita
        or emitir_receita_especial
        or emitir_histo
        or emitir_antibio
    )

    if algum_selecionado:
      st.write("---")
      st.subheader("2. Dados dos Documentos Selecionados")

      nome = st.text_input("Nome do paciente:").strip().upper()

      if emitir_histo or emitir_antibio:
        col_id1, col_id2 = st.columns(2)
        with col_id1:
          data_nascimento = st.text_input(
              "Data de nascimento (Ex: 12/05/1985):"
          ).strip()
        with col_id2:
          nome_mae = st.text_input("Nome da mãe ou pai:").strip().upper()
      else:
        data_nascimento = ""
        nome_mae = ""

      cid_final = ""
      incluir_cid_atestado = False
      afastamento = 0
      if emitir_atestado:
        col_p1, col_p2 = st.columns(2)
        with col_p1:
          afastamento = st.number_input(
              "Dias de afastamento:", min_value=0, value=15, step=1
          )
        with col_p2:
          st.write("\n")
          incluir_cid_atestado = st.checkbox(
              "Incluir CID-10 e autorização do paciente", value=True
          )

        if incluir_cid_atestado:
          df_cid = carregar_biblioteca_cid10()
          cid_sel = st.selectbox(
              "Pesquisar CID-10 (digite o código ou nome da doença):",
              options=[""] + list(df_cid["label"])
              if not df_cid.empty
              else [""],
          )
          cid_manual = (
              st.text_input("Ou digite o CID-10 manualmente:").strip().upper()
          )
          cid_final = cid_sel if cid_sel else cid_manual

      retorno = 0
      if emitir_retorno:
        retorno = st.number_input(
            "Dias até o retorno:", min_value=0, value=14, step=1
        )

      prescricao_especial = ""
      nome_medico_especial = ""
      if emitir_receita_especial:
        nome_medico_especial = st.text_input(
            "Nome do médico (Identificação do Emitente):"
        ).strip().upper()
        prescricao_especial = st.text_area(
            "Prescrição para Receituário Especial (Ex: Medicação, dose, via e"
            " posologia):"
        ).strip()

      material_exame = ""
      quadro_clinico = ""
      usa_antibiotico = "NÃO"
      antibiotico_nome_dias = ""
      if emitir_histo or emitir_antibio:
        material_exame = st.text_input("Material coletado:").strip().upper()
        quadro_clinico = st.text_area("Quadro clínico:").strip().upper()

        if emitir_antibio:
          col_ab1, col_ab2 = st.columns(2)
          with col_ab1:
            usa_antibiotico = st.selectbox(
                "Faz uso de antibiótico?", ["NÃO", "SIM"]
            )
          with col_ab2:
            if usa_antibiotico == "SIM":
              antibiotico_nome_dias = st.text_input(
                  "Antibiótico(s) e dias de tratamento:"
              ).strip().upper()

      st.write("\n")
      if st.button("📦 Gerar Documentos (.DOCX)", type="primary"):
        if not nome:
          st.error("Por favor, preencha o Nome do paciente.")
        elif emitir_receita_especial and not nome_medico_especial:
          st.error(
              "Por favor, preencha o Nome do médico para o Receituário Especial."
          )
        else:
          hoje_str = datetime.now().strftime("%d/%m/%Y")

          data_nasc_txt = (
              data_nascimento
              if data_nascimento
              else "_________________________"
          )
          nome_mae_txt = (
              nome_mae
              if nome_mae
              else "__________________________________________________"
          )
          cid_txt = (
              cid_final
              if cid_final
              else "__________________________________________________"
          )

          documentos_gerar = []

          if emitir_atestado:
            corpo_atestado = (
                f"Nome do paciente: {nome}\n\n"
                "Atesto, para os devidos fins, que o(a) paciente acima"
                " identificado(a) recebeu atendimento neste serviço no dia"
                " ____/____/______, e necessita de afastamento de suas"
                f" atividades por {afastamento} dias, a partir desta data.\n\n"
            )
            if incluir_cid_atestado:
              corpo_atestado += (
                  f"CID-10: {cid_txt}\n\n"
                  "CID autorizado pelo paciente.\n\n"
                  "_____________________________________\nAssinatura do"
                  " paciente ou responsável legal\n\n\n"
                  "_____________________________\nAssinatura do médico"
              )
            else:
              corpo_atestado += (
                  "_____________________________\nAssinatura do médico"
              )

            documentos_gerar.append(("ATESTADO MÉDICO", corpo_atestado))

          if emitir_retorno:
            corpo_retorno = (
                f"Nome do paciente: {nome}\n\n"
                "Retorno ao Ambulatório de Cirurgia Geral em"
                f" {retorno} dias para reavaliação pós-operatória.\n\n"
                "Orientações:\n"
                "• Manter curativo limpo e seco.\n"
                "• Utilizar as medicações prescritas.\n"
                "• Evitar esforços físicos até nova avaliação.\n"
                "• Manter alimentação e hidratação adequadas.\n"
                "• Procurar atendimento de urgência em caso de febre, dor,"
                " vômitos, vermelhidão intensa, secreção pela ferida,"
                " sangramento ou abertura dos pontos.\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append((
                "ENCAMINHAMENTO PARA RETORNO – AMBULATÓRIO DE CIRURGIA GERAL",
                corpo_retorno,
            ))

          if emitir_receita:
            corpo_receita = (
                f"Nome do paciente:"
                f" {nome}\n\n\n\n\n\n\n\n\n\n_____________________________\nAssinatura"
                " do médico"
            )
            documentos_gerar.append(("RECEITUÁRIO COMUM", corpo_receita))

          if emitir_histo:
            corpo_histo = (
                f"NOME DO PACIENTE: {nome}\n"
                f"DATA DE NASCIMENTO: {data_nasc_txt}\n"
                f"NOME DA MÃE OU PAI: {nome_mae_txt}\n\n"
                f"MATERIAL: {material_exame}\n"
                "ANÁLISE: Histopatológico\n"
                f"QUADRO CLÍNICO: {quadro_clinico}\n\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append(
                ("SOLICITAÇÃO DE HISTOPATOLÓGICO", corpo_histo)
            )

          if emitir_antibio:
            corpo_antibio = (
                f"NOME DO PACIENTE: {nome}\n"
                f"DATA DE NASCIMENTO: {data_nasc_txt}\n"
                f"NOME DA MÃE OU PAI: {nome_mae_txt}\n\n"
                f"MATERIAL: {material_exame}\n"
                "ANÁLISE: Antibiograma\n"
                f"QUADRO CLÍNICO: {quadro_clinico}\n"
                f"FAZ USO DE ANTIBIÓTICO? {usa_antibiotico}\n"
            )
            if usa_antibiotico == "SIM":
              corpo_antibio += (
                  f"ANTIBIÓTICO(S) E DIAS: {antibiotico_nome_dias}\n\n\n"
              )
            else:
              corpo_antibio += "\n\n\n"
            corpo_antibio += "_____________________________\nAssinatura do médico"
            documentos_gerar.append(
                ("SOLICITAÇÃO DE ANTIBIOGRAMA", corpo_antibio)
            )

          doc_word = Document()

          def obter_imagem(nome_base):
            for ext in [".png", ".jpg", ".jpeg"]:
              if os.path.exists(nome_base + ext):
                return nome_base + ext
            return None

          img_topo, img_rodape = obter_imagem("topo"), obter_imagem("rodape")

          for section in doc_word.sections:
            section.top_margin = Inches(1.8)
            section.bottom_margin = Inches(1.6)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            if img_topo:
              section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              section.header.paragraphs[0].add_run().add_picture(
                  img_topo, width=Inches(6.0)
              )
            if img_rodape:
              section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              section.footer.paragraphs[0].add_run().add_picture(
                  img_rodape, width=Inches(6.0)
              )

          for i, (titulo, corpo) in enumerate(documentos_gerar):
            p_titulo = doc_word.add_paragraph()
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_titulo.paragraph_format.space_after = Pt(18)
            run_tit = p_titulo.add_run(titulo)
            run_tit.font.name = "Arial"
            run_tit.font.size = Pt(13)
            run_tit.bold = True

            for linha in corpo.split("\n"):
              if linha.strip() == "":
                doc_word.add_paragraph().paragraph_format.space_after = Pt(4)
                continue
              p_linha = doc_word.add_paragraph()
              p_linha.paragraph_format.line_spacing = 1.2
              p_linha.paragraph_format.space_after = Pt(4)
              p_linha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

              run_linha = p_linha.add_run(linha)
              run_linha.font.name = "Arial"
              run_linha.font.size = Pt(11)

            p_data = doc_word.add_paragraph()
            p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_data.paragraph_format.space_before = Pt(20)
            run_data = p_data.add_run(f"{hoje_str}, Campo Maior - PI")
            run_data.font.name = "Arial"
            run_data.font.size = Pt(10)

            if i < len(documentos_gerar) - 1 or emitir_receita_especial:
              doc_word.add_page_break()

          if emitir_receita_especial:
            adicionar_receituario_especial(
                doc_word, nome_medico_especial, nome, prescricao_especial
            )

          bio = io.BytesIO()
          doc_word.save(bio)

          st.success("✨ Documentos gerados com sucesso!")
          st.download_button(
              label="📥 Baixar Documentos .DOCX",
              data=bio.getvalue(),
              file_name=(
                  f"Documentos_Cirurgia_Geral_{nome.lower().replace(' ', '_')}.docx"
              ),
              mime=(
                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              ),
          )

  # ---------------------------------------------------------
  # ORTOPEDIA
  # ---------------------------------------------------------
  elif st.session_state["pagina_atual"] == "ortopedia":
    if st.button("⬅️ Voltar ao Menu Principal"):
      st.session_state["pagina_atual"] = "menu"
      st.rerun()

    st.header("🦴 Emissor de Documentos - Ortopedia e Traumatologia")

    st.subheader("1. Selecione os Documentos a serem gerados")
    emitir_atestado = st.checkbox("Atestado Médico", value=False)
    emitir_retorno = st.checkbox(
        "Ficha de Retorno (Ambulatório)", value=False
    )
    emitir_receita = st.checkbox("Receituário Comum", value=False)
    emitir_receita_especial = st.checkbox(
        "Receituário de Controle Especial (2 Vias)", value=False
    )
    emitir_rx_retorno = st.checkbox(
        "Solicitação de Radiografia de Retorno", value=False
    )
    emitir_imobilizacao = st.checkbox(
        "Solicitação de Imobilização", value=False
    )
    emitir_fisio = st.checkbox("Solicitação de Fisioterapia", value=False)
    emitir_antibio = st.checkbox("Solicitação de Antibiograma", value=False)

    algum_selecionado = (
        emitir_atestado
        or emitir_retorno
        or emitir_receita
        or emitir_receita_especial
        or emitir_rx_retorno
        or emitir_imobilizacao
        or emitir_fisio
        or emitir_antibio
    )

    if algum_selecionado:
      st.write("---")
      st.subheader("2. Dados dos Documentos Selecionados")

      nome = st.text_input("Nome do paciente:").strip().upper()

      if emitir_antibio:
        col_id1, col_id2 = st.columns(2)
        with col_id1:
          data_nascimento = st.text_input(
              "Data de nascimento (Ex: 12/05/1985):"
          ).strip()
        with col_id2:
          nome_mae = st.text_input("Nome da mãe ou pai:").strip().upper()
      else:
        data_nascimento = ""
        nome_mae = ""

      cid_final = ""
      if emitir_atestado or emitir_rx_retorno or emitir_imobilizacao:
        df_cid = carregar_biblioteca_cid10()
        cid_sel = st.selectbox(
            "Pesquisar CID-10 (digite o código ou nome da doença):",
            options=[""] + list(df_cid["label"]) if not df_cid.empty else [""],
        )
        cid_manual = (
            st.text_input("Ou digite o CID-10 manualmente:").strip().upper()
        )
        cid_final = cid_sel if cid_sel else cid_manual

      incluir_cid_atestado = False
      afastamento = 0
      if emitir_atestado:
        col_p1, col_p2 = st.columns(2)
        with col_p1:
          afastamento = st.number_input(
              "Dias de afastamento:", min_value=0, value=45, step=1
          )
        with col_p2:
          st.write("\n")
          incluir_cid_atestado = st.checkbox(
              "Incluir CID-10 e autorização do paciente no atestado", value=True
          )

      retorno = 0
      if emitir_retorno:
        retorno = st.number_input(
            "Dias até o retorno:", min_value=0, value=14, step=1
        )

      prescricao_especial = ""
      nome_medico_especial = ""
      if emitir_receita_especial:
        nome_medico_especial = st.text_input(
            "Nome do médico (Identificação do Emitente):"
        ).strip().upper()
        prescricao_especial = st.text_area(
            "Prescrição para Receituário Especial (Ex: Medicação, dose, via e"
            " posologia):"
        ).strip()

      region_rx = ""
      incidencias_rx = ""
      hd_rx = ""
      if emitir_rx_retorno:
        col_rx1, col_rx2 = st.columns(2)
        with col_rx1:
          region_rx = st.text_input(
              "Região anatômica para Raio-X (Ex: JOELHO DIREITO):"
          ).strip().upper()
        with col_rx2:
          incidencias_rx = st.text_input(
              "Incidências (Ex: AP + PERFIL):"
          ).strip().upper()
        hd_rx = st.text_input(
            "Hipótese diagnóstica / CID (Raio-X):", value=cid_final
        ).strip().upper()

      tipo_imobilizacao = ""
      hd_imobilizacao = ""
      if emitir_imobilizacao:
        col_im1, col_im2 = st.columns(2)
        with col_im1:
          tipo_imobilizacao = st.text_input(
              "Tipo de imobilização (Ex: TALA GESSO MALARES E PALMAR):"
          ).strip().upper()
        with col_im2:
          hd_imobilizacao = st.text_input(
              "Hipótese diagnóstica (HD - Imobilização):", value=cid_final
          ).strip().upper()

      sessoes_fisio = 0
      quadro_clinico_fisio = ""
      if emitir_fisio:
        sessoes_fisio = st.number_input(
            "Quantidade de sessões de fisioterapia motora:",
            min_value=1,
            value=20,
            step=1,
        )
        quadro_clinico_fisio = st.text_area(
            "Quadro clínico para fisioterapia:"
        ).strip().upper()

      material_exame = ""
      quadro_clinico_antibiograma = ""
      usa_antibiotico = "NÃO"
      antibiotico_nome_dias = ""
      if emitir_antibio:
        material_exame = st.text_input(
            "Material coletado (Antibiograma):"
        ).strip().upper()
        quadro_clinico_antibiograma = st.text_area(
            "Quadro clínico (Antibiograma):"
        ).strip().upper()
        col_ab1, col_ab2 = st.columns(2)
        with col_ab1:
          usa_antibiotico = st.selectbox(
              "Faz uso de antibiótico?", ["NÃO", "SIM"]
          )
        with col_ab2:
          if usa_antibiotico == "SIM":
            antibiotico_nome_dias = st.text_input(
                "Antibiótico(s) e dias de tratamento:"
            ).strip().upper()

      st.write("\n")
      if st.button(
          "📦 Gerar Documentos Ortopédicos (.DOCX)", type="primary"
      ):
        if not nome:
          st.error("Por favor, preencha o Nome do paciente.")
        elif emitir_receita_especial and not nome_medico_especial:
          st.error(
              "Por favor, preencha o Nome do médico para o Receituário Especial."
          )
        else:
          hoje_str = datetime.now().strftime("%d/%m/%Y")

          data_nasc_txt = (
              data_nascimento
              if data_nascimento
              else "_________________________"
          )
          nome_mae_txt = (
              nome_mae
              if nome_mae
              else "__________________________________________________"
          )
          cid_txt = (
              cid_final
              if cid_final
              else "__________________________________________________"
          )

          documentos_gerar = []

          if emitir_atestado:
            corpo_atestado = (
                f"Nome do paciente: {nome}\n\n"
                "Atesto, para os devidos fins, que o(a) paciente acima"
                " identificado(a) recebeu atendimento neste serviço no dia"
                " ____/____/______, e necessita de afastamento de suas"
                f" atividades por {afastamento} dias, a partir desta data.\n\n"
            )
            if incluir_cid_atestado:
              corpo_atestado += (
                  f"CID-10: {cid_txt}\n\n"
                  "CID autorizado pelo paciente.\n\n"
                  "_____________________________________\nAssinatura do"
                  " paciente ou responsável legal\n\n\n"
                  "_____________________________\nAssinatura do médico"
              )
            else:
              corpo_atestado += (
                  "_____________________________\nAssinatura do médico"
              )

            documentos_gerar.append(("ATESTADO MÉDICO", corpo_atestado))

          if emitir_retorno:
            corpo_retorno = (
                f"Nome do paciente: {nome}\n\n"
                "Retorno ao Ambulatório de Ortopedia e Traumatologia em"
                f" {retorno} dias para reavaliação clínica e evolução do"
                " tratamento.\n\n"
                "Orientações:\n"
                "• Realizar raio-x de retorno, se houver.\n"
                "• Utilizar as medicações prescritas.\n"
                "• Procurar atendimento de urgência em caso de febre, dor,"
                " vermelhidão, secreção e/ou calor local, abertura dos pontos"
                " ou se exposição de placas/pinos/parafusos, se houver.\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append((
                "ENCAMINHAMENTO PARA RETORNO – AMBULATÓRIO DE ORTOPEDIA E"
                " TRAUMATOLOGIA",
                corpo_retorno,
            ))

          if emitir_receita:
            corpo_receita = (
                f"Nome do paciente:"
                f" {nome}\n\n\n\n\n\n\n\n\n\n_____________________________\nAssinatura"
                " do médico"
            )
            documentos_gerar.append(("RECEITUÁRIO COMUM", corpo_receita))

          if emitir_rx_retorno:
            rx_txt = f"RAIO-X DE {region_rx}"
            if incidencias_rx:
              rx_txt += f" ({incidencias_rx})"

            hd_rx_txt = (
                hd_rx
                if hd_rx
                else "__________________________________________________"
            )

            corpo_rx = (
                "SOLICITAÇÃO DE PROCEDIMENTO\n\n"
                f"Nome do Paciente: {nome}\n\n"
                "SOLICITO:\n\n"
                f"{rx_txt}\n\n"
                f"HD: {hd_rx_txt}\n\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append(("SOLICITAÇÃO DE EXAMES", corpo_rx))

          if emitir_imobilizacao:
            corpo_imobilizacao = (
                f"Nome do paciente: {nome}\n\n"
                "SOLICITO:\n\n"
                f"IMOBILIZAÇÃO TIPO: {tipo_imobilizacao}\n\n"
                f"HD: {hd_imobilizacao}\n\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append(
                ("SOLICITAÇÃO DE IMOBILIZAÇÃO", corpo_imobilizacao)
            )

          if emitir_fisio:
            corpo_fisio = (
                f"Nome do paciente: {nome}\n\n"
                "SOLICITO:\n\n"
                f"{sessoes_fisio} sessões de fisioterapia motora.\n\n"
                f"Quadro clínico: {quadro_clinico_fisio}\n\n\n\n"
                "_____________________________\nAssinatura do médico"
            )
            documentos_gerar.append(
                ("FICHA DE ENCAMINHAMENTO PARA FISIOTERAPIA", corpo_fisio)
            )

          if emitir_antibio:
            corpo_antibio = (
                f"NOME DO PACIENTE: {nome}\n"
                f"DATA DE NASCIMENTO: {data_nasc_txt}\n"
                f"NOME DA MÃE OU PAI: {nome_mae_txt}\n\n"
                f"MATERIAL: {material_exame}\n"
                "ANÁLISE: Antibiograma\n"
                f"QUADRO CLÍNICO: {quadro_clinico_antibiograma}\n"
                f"FAZ USO DE ANTIBIÓTICO? {usa_antibiotico}\n"
            )
            if usa_antibiotico == "SIM":
              corpo_antibio += (
                  f"ANTIBIÓTICO(S) E DIAS: {antibiotico_nome_dias}\n\n\n"
              )
            else:
              corpo_antibio += "\n\n\n"
            corpo_antibio += "_____________________________\nAssinatura do médico"
            documentos_gerar.append(
                ("SOLICITAÇÃO DE ANTIBIOGRAMA", corpo_antibio)
            )

          doc_word = Document()

          def obter_imagem(nome_base):
            for ext in [".png", ".jpg", ".jpeg"]:
              if os.path.exists(nome_base + ext):
                return nome_base + ext
            return None

          img_topo, img_rodape = obter_imagem("topo"), obter_imagem("rodape")

          for section in doc_word.sections:
            section.top_margin = Inches(1.8)
            section.bottom_margin = Inches(1.6)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            if img_topo:
              section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              section.header.paragraphs[0].add_run().add_picture(
                  img_topo, width=Inches(6.0)
              )
            if img_rodape:
              section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              section.footer.paragraphs[0].add_run().add_picture(
                  img_rodape, width=Inches(6.0)
              )

          for i, (titulo, corpo) in enumerate(documentos_gerar):
            p_titulo = doc_word.add_paragraph()
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_titulo.paragraph_format.space_after = Pt(18)
            run_tit = p_titulo.add_run(titulo)
            run_tit.font.name = "Arial"
            run_tit.font.size = Pt(13)
            run_tit.bold = True

            for linha in corpo.split("\n"):
              if linha.strip() == "":
                doc_word.add_paragraph().paragraph_format.space_after = Pt(4)
                continue
              p_linha = doc_word.add_paragraph()
              p_linha.paragraph_format.line_spacing = 1.2
              p_linha.paragraph_format.space_after = Pt(4)
              p_linha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

              run_linha = p_linha.add_run(linha)
              run_linha.font.name = "Arial"
              run_linha.font.size = Pt(11)

            p_data = doc_word.add_paragraph()
            p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_data.paragraph_format.space_before = Pt(20)
            run_data = p_data.add_run(f"{hoje_str}, Campo Maior - PI")
            run_data.font.name = "Arial"
            run_data.font.size = Pt(10)

            if i < len(documentos_gerar) - 1 or emitir_receita_especial:
              doc_word.add_page_break()

          if emitir_receita_especial:
            adicionar_receituario_especial(
                doc_word, nome_medico_especial, nome, prescricao_especial
            )

          bio = io.BytesIO()
          doc_word.save(bio)

          st.success("✨ Documentos de Ortopedia gerados com sucesso!")
          st.download_button(
              label="📥 Baixar Documentos .DOCX",
              data=bio.getvalue(),
              file_name=(
                  f"Documentos_Ortopedia_{nome.lower().replace(' ', '_')}.docx"
              ),
              mime=(
                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              ),
          )

elif senha != "":
  st.error("Chave de acesso inválida.")
